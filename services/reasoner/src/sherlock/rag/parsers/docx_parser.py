"""Parser for DOCX files using python-docx."""
from __future__ import annotations

import io

import docx

from sherlock.rag.domain.models import ParsedDocument


class DocxParser:
    def parse(self, data: bytes) -> ParsedDocument:
        doc = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        return ParsedDocument(text=text, metadata={"type": "docx"})
